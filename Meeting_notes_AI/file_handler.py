"""
File handling service for extracting text from uploaded documents (.txt, .pdf, .docx).
"""

import io
from typing import Optional
import PyPDF2
import docx


class FileExtractionError(Exception):
    """Custom exception raised during file reading failures."""
    pass


def extract_text_from_txt(file_bytes: bytes) -> str:
    """
    Extracts text from plain text file bytes.
    
    Args:
        file_bytes (bytes): Binary file data.
        
    Returns:
        str: Decoded string.
    """
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from a PDF file buffer.
    
    Args:
        file_bytes (bytes): PDF binary data.
        
    Returns:
        str: Extracted text.
    """
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PyPDF2.PdfReader(pdf_file)
        extracted_text = []
        for idx, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                extracted_text.append(text)
        return "\n".join(extracted_text)
    except Exception as e:
        raise FileExtractionError(f"Error parsing PDF file: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from a Microsoft Word (.docx) file buffer.
    
    Args:
        file_bytes (bytes): DOCX binary data.
        
    Returns:
        str: Extracted text from paragraphs and tables.
    """
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        full_text = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    full_text.append(row_text)
        return "\n".join(full_text)
    except Exception as e:
        raise FileExtractionError(f"Error parsing DOCX file: {str(e)}")


def process_uploaded_file(uploaded_file) -> Optional[str]:
    """
    Dispatcher method to handle uploaded file formats.
    
    Args:
        uploaded_file: Streamlit UploadedFile instance.
        
    Returns:
        Optional[str]: Extracted text payload.
    """
    if uploaded_file is None:
        return None

    filename = uploaded_file.name.lower()
    content = uploaded_file.read()

    if filename.endswith(".txt"):
        return extract_text_from_txt(content)
    elif filename.endswith(".pdf"):
        return extract_text_from_pdf(content)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(content)
    else:
        raise FileExtractionError("Unsupported file extension. Please upload .txt, .pdf, or .docx.")